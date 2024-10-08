import boto3
import os
import io
import email
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.utils import simpleSplit

s3 = boto3.client('s3')

def extract_doc_data(doc_binary):
    doc_buffer = io.BytesIO(doc_binary)
    try:
        doc = Document(doc_buffer)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append({
                    'type': 'text',
                    'text': paragraph.text.strip()
                })
        
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    content.append({
                        'type': 'text',
                        'text': row_text
                    })
        
        return content
    except Exception as e:
        raise Exception(f"Failed to read Word document: {str(e)}")

def create_pdf_from_doc(content):
    buffer = io.BytesIO()
    page_width, page_height = landscape(letter)
    c = canvas.Canvas(buffer, pagesize=landscape(letter))
    
    margin = 50
    y = page_height - margin
    available_width = page_width - (2 * margin)
    
    def write_text_block(text, y_pos):
        if not text.strip():
            return y_pos
        
        lines = simpleSplit(text, "Helvetica", 10, available_width)
        for line in lines:
            if y_pos < margin + 20:
                c.showPage()
                y_pos = page_height - margin
                c.setFont("Helvetica", 10)
            c.drawString(margin, y_pos, line)
            y_pos -= 14
        return y_pos - 10
    
    try:
        for item in content:
            if y < margin + 50:
                c.showPage()
                y = page_height - margin
            
            c.setFont("Helvetica", 10)
            y = write_text_block(item['text'], y)
        
        c.save()
        return buffer.getvalue()
    except Exception as e:
        raise Exception(f"Error creating PDF: {str(e)}")

def handler(event, context):
    print(f"Processing Word Document attachment...")
    bucket_name = os.environ['BUCKET_NAME']
    message_id = event['messageId']
    attachment_filename = event['filename']
    try:
        obj = s3.get_object(Bucket=bucket_name, Key=message_id)
        email_content = obj['Body'].read().decode('utf-8')
        
        doc_data = None
        email_message = email.message_from_string(email_content)
        for part in email_message.walk():
            if part.get_content_maintype() == 'application' and part.get_filename() == attachment_filename:
                doc_data = part.get_payload(decode=True)
                break
        
        if not doc_data:
            return {
                'statusCode': 404,
                'status': 'error',
                'body': f'No Word document attachment found'
            }
        
        doc_content = extract_doc_data(doc_data)
        pdf_data = create_pdf_from_doc(doc_content)
        
        original_filename = os.path.splitext(attachment_filename)[0]
        pdf_key = f'invoices/{message_id}/{original_filename}.pdf'
        s3.put_object(
            Bucket=bucket_name,
            Key=pdf_key,
            Body=pdf_data,
            ContentType='application/pdf'
        )
        
        return {
            'statusCode': 200,
            'status': 'success',
            'pdfKey': pdf_key
        }
        
    except Exception as e:
        print(f"Error processing Word document: {str(e)}")
        return {
            'statusCode': 500,
            'status': 'error',
            'error': str(e)
        }